import os
import uuid
import re
import requests
import hmac
import hashlib
import io
from datetime import datetime, timezone, timedelta
from flask import Blueprint, render_template, redirect, url_for, flash, request, current_app, jsonify, send_file, Response
from werkzeug.utils import secure_filename, safe_join
from .models import User, File, Chunk, UploadUsage, Folder, Tag, file_tags, BillingEvent
from . import db, limiter
from flask_login import login_user, login_required, logout_user, current_user
from flask_limiter.util import get_remote_address
from .document_processor import extract_text
from .faiss_index import FaissIndex
from .embeddings import EmbeddingGenerator

main = Blueprint("main", __name__)

LS_API_BASE = "https://api.lemonsqueezy.com/v1"

# -------------------------------------------------------
# Plan limits
# -------------------------------------------------------
FREE_LIMITS = {
    "max_uploads_per_month": 10,
    "max_queries_per_month": 100,
    "storage_mb": 100,
    "max_file_size_bytes": 5 * 1024 * 1024,
    "allowed_extensions": {".pdf"},
}

PRO_LIMITS = {
    "max_uploads_per_month": None,  # unlimited (fair-use)
    "max_queries_per_month": 500,
    "storage_mb": 2048,
    "max_file_size_bytes": 100 * 1024 * 1024,
    "allowed_extensions": {".txt", ".md", ".pdf", ".docx", ".pptx",
                           ".xlsx", ".csv", ".html", ".json"},
}

def get_limits(user):
    return PRO_LIMITS if user.is_pilot else FREE_LIMITS

def get_current_month_key():
    return datetime.utcnow().strftime("%Y-%m")

def get_monthly_upload_count(user_id):
    month_key = get_current_month_key()
    usage = UploadUsage.query.filter_by(user_id=user_id, month_key=month_key).first()
    return usage.upload_count if usage else 0

def increment_monthly_uploads(user_id):
    month_key = get_current_month_key()
    usage = UploadUsage.query.filter_by(user_id=user_id, month_key=month_key).first()
    if not usage:
        usage = UploadUsage(user_id=user_id, month_key=month_key, upload_count=0)
        db.session.add(usage)
    usage.upload_count += 1
    db.session.commit()

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".pptx",
                      ".xlsx", ".csv", ".html", ".json"}

def allowed_file(filename, user):
    ext = os.path.splitext(filename)[1].lower()
    return ext in get_limits(user)["allowed_extensions"]

def check_storage_space(user_id, required_space=0):
    try:
        user = User.query.get(user_id)
        limits = get_limits(user)
        limit_mb = limits["storage_mb"]

        user_folder = os.path.join(current_app.config["UPLOAD_FOLDER"], str(user_id))
        if not os.path.exists(user_folder):
            return True, limit_mb, 0.0

        total_size = 0
        for dirpath, dirnames, filenames in os.walk(user_folder):
            for fname in filenames:
                fp = os.path.join(dirpath, fname)
                if os.path.exists(fp):
                    total_size += os.path.getsize(fp)

        used_mb = total_size / (1024 * 1024)
        available_mb = limit_mb - used_mb
        has_space = available_mb >= (required_space / (1024 * 1024))
        return has_space, available_mb, used_mb
    except Exception:
        return True, 500, 0.0

def check_system_load():
    try:
        import psutil
        cpu_percent = psutil.cpu_percent(interval=1)
        memory_percent = psutil.virtual_memory().percent
        return cpu_percent > 90 or memory_percent > 90, max(cpu_percent, memory_percent)
    except Exception:
        return False, 0.0


@main.route("/health")
def health():
    return {"status": "ok"}, 200


# -----------------------
# Auth Routes
# -----------------------
from flask_mail import Message
from app import mail, csrf


@main.route("/")
def home():
    return render_template("index.html", user=current_user)


@main.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username")
        email = request.form.get("email")
        password = request.form.get("password")

        if User.query.filter_by(username=username).first():
            flash("Username already exists.", "danger")
            return redirect(url_for("main.register"))
        if User.query.filter_by(email=email).first():
            flash("Email already exists.", "danger")
            return redirect(url_for("main.register"))

        new_user = User(username=username, email=email)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()

        # Email optionnel — ne bloque pas l'inscription
        try:
            mail_server = os.getenv("MAIL_SERVER")
            if mail_server:
                msg = Message(
                    subject="Welcome to Doxium 🚀",
                    recipients=[new_user.email]
                )
                msg.body = f"Hi {new_user.username},\n\nWelcome to Doxium!\n\nBest,\nThe Doxium Team"
                from threading import Thread
                def send_async(app, message):
                    with app.app_context():
                        try:
                            mail.send(message)
                        except Exception:
                            pass
                Thread(target=send_async, args=(current_app._get_current_object(), msg)).start()
        except Exception:
            pass

        flash("Registration successful! Please log in.", "success")
        return redirect(url_for("main.login"))

    return render_template("register.html")


from flask_limiter.errors import RateLimitExceeded

failed_login_limit = limiter.limit("5 per minute", key_func=get_remote_address)


@main.route("/login", methods=["GET", "POST"])
@failed_login_limit
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        user = User.query.filter_by(email=email).first()

        if user and user.check_password(password):
            login_user(user)
            if user.is_pilot:
                flash("Welcome back!", "success")
            else:
                flash("Welcome back! You're on the Free plan.", "info")
            return redirect(url_for("main.dashboard"))

        flash("Invalid email or password.", "danger")

    return render_template("login.html")


@main.route("/logout")
@login_required
def logout():
    logout_user()
    flash("Logged out successfully.", "success")
    return redirect(url_for("main.login"))


@main.route("/create-checkout-session", methods=["POST"])
@login_required
@csrf.exempt
def create_checkout_session():
    if current_user.is_pilot:
        flash("You're already on Pro.", "info")
        return redirect(url_for("main.dashboard"))
    try:
        checkout_url = _create_ls_checkout_url(
            user_id=str(current_user.id),
            redirect_url=url_for("main.payment_success", _external=True),
        )
        return redirect(checkout_url, code=303)
    except Exception as e:
        current_app.logger.exception("Failed to create Lemon Squeezy checkout")
        return jsonify(error=str(e)), 400


@main.route("/create-checkout", methods=["POST"])
@login_required
@csrf.exempt
def create_checkout():
    if current_user.is_pilot:
        return jsonify({"error": "already_pilot"}), 400

    data = request.get_json(silent=True) or {}
    redirect_url = data.get("redirect_url") or url_for("main.payment_success", _external=True)

    try:
        checkout_url = _create_ls_checkout_url(
            user_id=str(current_user.id),
            redirect_url=redirect_url,
        )
        return jsonify({"checkout_url": checkout_url})
    except Exception as e:
        current_app.logger.exception("Failed to create Lemon Squeezy checkout")
        return jsonify({"error": "failed_to_create_checkout"}), 400

import zipfile
import io

@main.route("/download-all-files")
@login_required
def download_all_files():
    files = File.query.filter_by(user_id=current_user.id).all()
    if not files:
        flash("No files to download.", "info")
        return redirect(url_for("main.dashboard"))

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in files:
            if os.path.exists(file.path):
                zf.write(file.path, arcname=file.filename)

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"doxium_files_{current_user.username}.zip"
    )

@main.route("/payment-success")
@login_required
def payment_success():
    if current_user.is_pilot:
        flash("You're all set! Pro access is active.", "success")
        return redirect(url_for("main.dashboard"))
    return render_template("success.html", user=current_user)


@main.route("/api/me-status")
@login_required
def me_status():
    return jsonify({
        "is_pilot": bool(current_user.is_pilot),
        "subscription_status": current_user.subscription_status,
    })


@main.route("/pricing")
def pricing():
    return render_template("pricing.html", user=current_user)


@main.route("/webhook", methods=["POST"])
@main.route("/webhook/lemonsqueezy", methods=["POST"])
@csrf.exempt
def lemonsqueezy_webhook():
    import hmac as hmac_module

    webhook_secret = os.getenv("LEMON_SQUEEZY_WEBHOOK_SECRET")
    if not webhook_secret:
        current_app.logger.error("WEBHOOK ERROR: secret not configured")
        return jsonify(error="Webhook secret not configured"), 500

    payload = request.get_data()
    current_app.logger.info("WEBHOOK RECEIVED: %s bytes", len(payload))

    # Get signature from any possible header
    sig_header = (
        request.headers.get("X-Signature")
        or request.headers.get("X-Signature-256")
        or request.headers.get("X-Lemon-Signature")
    )

    if not sig_header:
        current_app.logger.error("WEBHOOK ERROR: no signature header found")
        return jsonify(error="Missing signature"), 400

    # Verify HMAC
    computed = hmac_module.new(
        webhook_secret.encode("utf-8"),
        payload,
        hashlib.sha256
    ).hexdigest()

    sig_value = sig_header.split("=", 1)[-1] if "=" in sig_header else sig_header

    if not hmac_module.compare_digest(computed, sig_value):
        current_app.logger.error("WEBHOOK ERROR: invalid signature")
        return jsonify(error="Invalid signature"), 400

    # Parse event
    event_name = request.headers.get("X-Event-Name", "")
    event = request.get_json(silent=True) or {}
    data = event.get("data") or {}
    attrs = data.get("attributes") or {}
    meta = event.get("meta") or {}

    if not event_name:
        event_name = meta.get("event_name", "")
    event_name = event_name.lower().strip()

    current_app.logger.info("WEBHOOK EVENT: '%s'", event_name)

    # Get user_id from custom_data
    custom = meta.get("custom_data") or attrs.get("custom_data") or {}
    user_id = custom.get("user_id")
    current_app.logger.info("WEBHOOK custom_data: %s", custom)

    # Find user
    user = None
    if user_id:
        try:
            user = User.query.get(int(user_id))
            current_app.logger.info("WEBHOOK user found by ID: %s", user_id)
        except Exception:
            pass

    if not user:
        email = (
            attrs.get("user_email")
            or attrs.get("customer_email")
            or attrs.get("email")
        )
        if email:
            user = User.query.filter_by(email=email).first()
            current_app.logger.info("WEBHOOK user found by email: %s", email)

    if not user:
        current_app.logger.warning("WEBHOOK: no user found, skipping")
        return jsonify(success=True)

    current_app.logger.info("WEBHOOK processing for user ID=%s", user.id)

    # Log billing event
    try:
        billing_event = BillingEvent(
            user_id=user.id,
            event_type=event_name,
            ls_order_id=str(data.get("id", "")),
            raw_payload=event,
        )
        db.session.add(billing_event)
        db.session.commit()
    except Exception as e:
        current_app.logger.error("WEBHOOK billing log error: %s", e)
        db.session.rollback()

    # Handle events
    if event_name == "order_created":
        _set_user_pilot(user, is_pilot=True, status="active",
                        purchased_at=datetime.utcnow())
        current_app.logger.info("WEBHOOK SUCCESS: user %s activated as pilot", user.id)

    elif event_name == "order_refunded":
        _set_user_pilot(user, is_pilot=False, status="refunded")
        current_app.logger.info("WEBHOOK: user %s refunded", user.id)

    elif event_name.startswith("subscription_"):
        sub_id = attrs.get("id") or data.get("id")
        ends_at = attrs.get("ends_at") or attrs.get("renews_at")
        portal_url = (attrs.get("urls") or {}).get("customer_portal")

        if event_name in (
            "subscription_created",
            "subscription_updated",
            "subscription_resumed",
            "subscription_payment_success",
        ):
            _set_user_pilot(
                user, is_pilot=True, status="active",
                subscription_id=sub_id,
                expires_at=_parse_ls_datetime(ends_at),
                portal_url=portal_url,
            )
            current_app.logger.info("WEBHOOK SUCCESS: user %s activated via %s",
                                     user.id, event_name)

        elif event_name in ("subscription_cancelled", "subscription_canceled"):
            _set_user_pilot(
                user, is_pilot=True, status="cancelled",
                subscription_id=sub_id,
                expires_at=_parse_ls_datetime(ends_at),
                portal_url=portal_url,
            )

        elif event_name == "subscription_expired":
            _set_user_pilot(
                user, is_pilot=False, status="expired",
                subscription_id=sub_id,
                expires_at=_parse_ls_datetime(ends_at),
            )

    return jsonify(success=True)


def _ls_headers():
    api_key = os.getenv("LEMON_SQUEEZY_API_KEY")
    if not api_key:
        raise ValueError("LEMON_SQUEEZY_API_KEY not configured")
    return {
        "Authorization": f"Bearer {api_key}",
        "Accept": "application/vnd.api+json",
        "Content-Type": "application/vnd.api+json",
    }


def _create_ls_checkout_url(user_id: str, redirect_url: str) -> str:
    store_id = os.getenv("LEMON_SQUEEZY_STORE_ID")
    variant_id = os.getenv("LEMON_SQUEEZY_VARIANT_ID")
    if not store_id or not variant_id:
        raise ValueError("LEMON_SQUEEZY_STORE_ID/LEMON_SQUEEZY_VARIANT_ID not configured")

    payload = {
        "data": {
            "type": "checkouts",
            "attributes": {
                "product_options": {
                    "redirect_url": url_for("main.payment_success", _external=True),
                    "enabled_variants": [int(variant_id)],
                },
                "checkout_data": {
                    "custom": {
                        "user_id": str(user_id)   # ← CRITIQUE
                    }
                },
            },
            "relationships": {
                "store": {"data": {"type": "stores", "id": str(store_id)}},
                "variant": {"data": {"type": "variants", "id": str(variant_id)}},
            },
        }
    }

    resp = requests.post(
        f"{LS_API_BASE}/checkouts",
        headers=_ls_headers(),
        json=payload,
        timeout=15
    )
    if resp.status_code not in (200, 201):
        raise RuntimeError(f"Lemon Squeezy checkout failed: {resp.text}")

    attrs = (resp.json().get("data") or {}).get("attributes", {})
    checkout_url = attrs.get("url") or attrs.get("checkout_url")
    if not checkout_url:
        raise RuntimeError("No checkout URL returned from Lemon Squeezy")
    return checkout_url


def _parse_ls_datetime(value):
    if not value:
        return None
    try:
        if value.endswith("Z"):
            value = value.replace("Z", "+00:00")
        return datetime.fromisoformat(value)
    except Exception:
        return None
from datetime import datetime

def _set_user_pilot(user, is_pilot, status=None, subscription_id=None, expires_at=None, purchased_at=None, portal_url=None):
    user.is_pilot = bool(is_pilot)
    if status:
        user.subscription_status = status
    if subscription_id:
        user.lemonsqueezy_subscription_id = str(subscription_id)
    if expires_at:
        user.subscription_expires_at = expires_at
    if purchased_at:
        user.pilot_purchased_at = purchased_at
    if portal_url:
        user.ls_customer_portal_url = portal_url
    if status in ("cancelled", "expired", "refunded") and not user.subscription_cancelled_at:
        user.subscription_cancelled_at = datetime.utcnow()
    if is_pilot:  # re-subscribing clears it
        user.subscription_cancelled_at = None
    db.session.commit()

# -----------------------
# Dashboard + File Management
# -----------------------
@main.route("/dashboard")
@login_required
def dashboard():
    files_with_content = []
    for file in current_user.files:
        files_with_content.append({
            "filename": file.filename,
            "size": file.size,
            "id": file.id,
            "processed": file.processed,
            "folder_id": file.folder_id,
            "tags": [{"id": t.id, "name": t.name, "color": t.color} for t in file.tags],
        })
    limits = get_limits(current_user)
    monthly_uploads = get_monthly_upload_count(current_user.id) if limits["max_uploads_per_month"] is not None else None
    total_size = sum(f.size for f in current_user.files)
    used_mb = total_size / (1024 * 1024)
    available_mb = limits["storage_mb"] - used_mb
    folders = Folder.query.filter_by(user_id=current_user.id).all()
    tags = Tag.query.filter_by(user_id=current_user.id).all()
    return render_template(
        "dashboard.html",
        user=current_user,
        files=files_with_content,
        is_pilot=current_user.is_pilot,
        limits=limits,
        monthly_uploads=monthly_uploads,
        storage_used_mb=used_mb,
        storage_available_mb=available_mb,
        folders=folders,
        tags=tags,
    )

@main.route("/upload", methods=["POST"])
@login_required
def upload():
    is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    def fail(msg):
        if is_ajax:
            return jsonify(success=False, error=msg)
        flash(msg, "danger")
        return redirect(url_for("main.dashboard"))

    limits = get_limits(current_user)

    try:
        is_overloaded, load_pct = check_system_load()
        if is_overloaded:
            return fail(f"System overloaded ({load_pct:.1f}%). Try again in a few minutes.")

        if "file" not in request.files:
            return fail("No file part.")

        f = request.files["file"]
        if not f.filename:
            return fail("No file selected.")

        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in limits["allowed_extensions"]:
            allowed = ", ".join(sorted(limits["allowed_extensions"]))
            return fail(f"File type not allowed on your plan. Allowed: {allowed}")

        f.seek(0, os.SEEK_END)
        size = f.tell()
        f.seek(0)

        if size == 0:
            return fail("File is empty.")

        if size > limits["max_file_size_bytes"]:
            max_mb = limits["max_file_size_bytes"] // (1024 * 1024)
            return fail(f"File too large. Max size on your plan is {max_mb} MB.")

        if limits["max_uploads_per_month"] is not None:
            monthly_count = get_monthly_upload_count(current_user.id)
            if monthly_count >= limits["max_uploads_per_month"]:
                return fail(f"You have reached the {limits['max_uploads_per_month']}-upload monthly limit on your plan.")

        has_space, available_mb, used_mb = check_storage_space(current_user.id, size)
        if not has_space:
            return fail(f"Storage limit reached. Used {used_mb:.1f} MB of {limits['storage_mb']} MB.")

        unique_filename = f"{uuid.uuid4().hex}{ext}"
        user_folder = os.path.join(current_app.config["UPLOAD_FOLDER"], str(current_user.id))
        os.makedirs(user_folder, exist_ok=True)
        filepath = os.path.join(user_folder, unique_filename)

        f.save(filepath)
        if not os.path.exists(filepath):
            raise Exception("Failed to save file to disk.")
        
        with open(filepath, 'rb') as file_obj:
            file_content = file_obj.read()
        
        new_file = File(
            filename=f.filename,
            path=filepath,
            size=size,
            user_id=current_user.id,
            content=file_content
        )
        db.session.add(new_file)
        db.session.commit()

        if limits["max_uploads_per_month"] is not None:
            increment_monthly_uploads(current_user.id)

        if is_ajax:
            return jsonify(success=True, message="File uploaded successfully.")
        flash("File uploaded successfully.", "success")
        return redirect(url_for("main.dashboard"))

    except Exception as e:
        if "filepath" in locals() and os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception:
                pass
        return fail(f"Upload failed: {str(e)}")


# -----------------------
# Process / Delete / View Routes
# -----------------------
@main.route("/process/<int:file_id>", methods=["POST"])
@login_required
def process_file_route(file_id):
    from .tasks import process_file_task

    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()

    try:
        process_file_task.delay(file.id, current_user.id, current_app.config["UPLOAD_FOLDER"])
        flash("File is being processed in the background.", "info")
    except Exception as e:
        flash(f"Error sending file to background processing: {str(e)}", "danger")

    return redirect(url_for("main.dashboard"))


@main.route("/delete/<int:file_id>", methods=["POST"])
@login_required
def delete_file(file_id):
    from .tasks import rebuild_index_task

    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    was_processed = file.processed

    for chunk in file.chunks:
        db.session.delete(chunk)
    if os.path.exists(file.path):
        os.remove(file.path)
    db.session.delete(file)
    db.session.commit()

    if was_processed:
        try:
            rebuild_index_task.delay(current_user.id)
            flash("File deleted. FAISS index is rebuilding.", "info")
        except Exception:
            flash("File deleted but FAISS index rebuild failed to start.", "warning")
    else:
        flash("File deleted successfully.", "success")

    return redirect(url_for("main.dashboard"))


@main.route("/view-file/<path:filepath>")
@login_required
def view_file(filepath: str):
    user_folder = safe_join(current_app.config["UPLOAD_FOLDER"], str(current_user.id))
    full_path = safe_join(user_folder, filepath)
    if os.path.exists(full_path) and os.path.isfile(full_path):
        if not any(full_path.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
            flash("File type not allowed.")
            return redirect(url_for("main.dashboard"))
        return send_file(full_path, as_attachment=True, download_name=os.path.basename(filepath))
    flash("File not found.")
    return redirect(url_for("main.dashboard"))


@main.route("/preview/<int:file_id>")
@login_required
def preview_file(file_id):
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    if os.path.exists(file.path):
        content = extract_text(file.path)
        return jsonify(success=True, content=content[:50000])
    return jsonify(success=True, content="No content available")


# -----------------------
# Query Routes
# -----------------------
@main.route("/query", methods=["POST"])
@limiter.limit("2 per minute")
@login_required
def query_documents():
    limits = get_limits(current_user)

    query_text = request.form.get("query", "").strip()
    if not query_text:
        flash("Please enter a query.", "danger")
        return redirect(url_for("main.dashboard"))

    if len(query_text) > 1000:
        flash("Query too long. Max 1000 characters.", "warning")
        return redirect(url_for("main.dashboard"))

    file_ids_raw = request.form.getlist("file_ids")
    file_ids = [int(fid) for fid in file_ids_raw if fid.isdigit()] or None

    from .tasks import generate_query_answer
    task = generate_query_answer.delay(current_user.id, query_text, file_ids)
    flash("Generating answer… This may take 10–30 seconds.", "info")
    return redirect(url_for("main.query_status", task_id=task.id))


@main.route("/query/status/<task_id>")
@login_required
def query_status(task_id):
    from celery.result import AsyncResult
    from app.celery_app import celery

    result = AsyncResult(task_id, app=celery)

    if result.ready():
        if result.successful():
            data = result.get()
            if "error" in data:
                flash(data["error"], "danger")
                return redirect(url_for("main.dashboard"))
            return render_template(
                "query_results.html",
                user=current_user,
                query=data["query"],
                answer=data["answer"],
                supporting_chunks=data["chunks"],
                task_id=task_id,
            )
        else:
            flash("Query generation failed. Please try again.", "danger")
            return redirect(url_for("main.dashboard"))

    return render_template("query_processing.html", task_id=task_id)


# -----------------------
# Folder & Tag Routes
# -----------------------
@main.route("/folders/create", methods=["POST"])
@login_required
def create_folder():
    name = request.form.get("name", "").strip()
    if not name:
        flash("Folder name is required.", "danger")
        return redirect(url_for("main.dashboard"))
    parent_id = request.form.get("parent_id")
    folder = Folder(name=name, user_id=current_user.id, parent_id=int(parent_id) if parent_id else None)
    db.session.add(folder)
    db.session.commit()
    flash(f"Folder '{name}' created.", "success")
    return redirect(url_for("main.dashboard"))


@main.route("/folders/<int:folder_id>/rename", methods=["POST"])
@login_required
def rename_folder(folder_id):
    folder = Folder.query.filter_by(id=folder_id, user_id=current_user.id).first_or_404()
    name = request.form.get("name", "").strip()
    if not name:
        flash("Folder name is required.", "danger")
        return redirect(url_for("main.dashboard"))
    folder.name = name
    db.session.commit()
    flash("Folder renamed.", "success")
    return redirect(url_for("main.dashboard"))


@main.route("/folders/<int:folder_id>/delete", methods=["POST"])
@login_required
def delete_folder(folder_id):
    folder = Folder.query.filter_by(id=folder_id, user_id=current_user.id).first_or_404()
    for f in folder.files:
        f.folder_id = None
    for child in folder.children:
        child.parent_id = folder.parent_id
    db.session.delete(folder)
    db.session.commit()
    flash("Folder deleted. Files moved to root.", "info")
    return redirect(url_for("main.dashboard"))


@main.route("/files/<int:file_id>/move", methods=["POST"])
@login_required
def move_file(file_id):
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    folder_id = request.form.get("folder_id")
    file.folder_id = int(folder_id) if folder_id else None
    db.session.commit()
    flash("File moved.", "success")
    return redirect(url_for("main.dashboard"))


@main.route("/tags/create", methods=["POST"])
@login_required
def create_tag():
    name = request.form.get("name", "").strip()
    color = request.form.get("color", "#6366f1")
    if not name:
        flash("Tag name is required.", "danger")
        return redirect(url_for("main.dashboard"))
    existing = Tag.query.filter_by(name=name, user_id=current_user.id).first()
    if existing:
        flash("Tag already exists.", "warning")
        return redirect(url_for("main.dashboard"))
    tag = Tag(name=name, user_id=current_user.id, color=color)
    db.session.add(tag)
    db.session.commit()
    flash(f"Tag '{name}' created.", "success")
    return redirect(url_for("main.dashboard"))


@main.route("/tags/<int:tag_id>/delete", methods=["POST"])
@login_required
def delete_tag(tag_id):
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    db.session.delete(tag)
    db.session.commit()
    flash("Tag deleted.", "success")
    return redirect(url_for("main.dashboard"))


@main.route("/files/<int:file_id>/tag", methods=["POST"])
@login_required
def tag_file(file_id):
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    tag_id = request.form.get("tag_id")
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    if tag not in file.tags:
        file.tags.append(tag)
        db.session.commit()
    flash("Tag added.", "success")
    return redirect(url_for("main.dashboard"))


@main.route("/files/<int:file_id>/untag", methods=["POST"])
@login_required
def untag_file(file_id):
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    tag_id = request.form.get("tag_id")
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    if tag in file.tags:
        file.tags.remove(tag)
        db.session.commit()
    flash("Tag removed.", "success")
    return redirect(url_for("main.dashboard"))


# -----------------------
# Study Feature Routes
# -----------------------
@main.route("/document/<int:file_id>/summarize", methods=["POST"])
@login_required
def summarize_document(file_id):
    if not current_user.is_pilot:
        flash("Upgrade to Pro for AI summaries.", "warning")
        return redirect(url_for("main.pricing"))
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    from .tasks import generate_summary
    task = generate_summary.delay(current_user.id, file_id)
    return redirect(url_for("main.study_status", task_id=task.id, study_type="summary"))


@main.route("/document/<int:file_id>/flashcards", methods=["POST"])
@login_required
def flashcards_document(file_id):
    if not current_user.is_pilot:
        flash("Upgrade to Pro for flashcards.", "warning")
        return redirect(url_for("main.pricing"))
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    from .tasks import generate_flashcards
    task = generate_flashcards.delay(current_user.id, file_id)
    return redirect(url_for("main.study_status", task_id=task.id, study_type="flashcards"))


@main.route("/document/<int:file_id>/quiz", methods=["POST"])
@login_required
def quiz_document(file_id):
    if not current_user.is_pilot:
        flash("Upgrade to Pro for quizzes.", "warning")
        return redirect(url_for("main.pricing"))
    file = File.query.filter_by(id=file_id, user_id=current_user.id).first_or_404()
    from .tasks import generate_quiz
    task = generate_quiz.delay(current_user.id, file_id)
    return redirect(url_for("main.study_status", task_id=task.id, study_type="quiz"))


@main.route("/study/status/<task_id>")
@login_required
def study_status(task_id):
    from celery.result import AsyncResult
    from app.celery_app import celery

    result = AsyncResult(task_id, app=celery)

    if result.ready():
        if result.successful():
            data = result.get()
            if "error" in data:
                flash(data["error"], "danger")
                return redirect(url_for("main.dashboard"))
            return render_template("study_results.html", user=current_user, data=data)
        else:
            flash("Study generation failed. Please try again.", "danger")
            return redirect(url_for("main.dashboard"))

    return render_template("query_processing.html", task_id=task_id)


# -----------------------
# Export Routes
# -----------------------
@main.route("/export/<fmt>/<task_id>")
@login_required
def export_query(fmt, task_id):
    from celery.result import AsyncResult
    from app.celery_app import celery

    result = AsyncResult(task_id, app=celery)
    if not result.ready() or not result.successful():
        flash("Result not ready or expired.", "danger")
        return redirect(url_for("main.dashboard"))

    data = result.get()
    answer = data.get("answer", "")
    query = data.get("query", "")
    chunks = data.get("chunks", [])

    sources_text = "\n\n---\nSources:\n"
    for c in chunks:
        sources_text += f"[{c.get('source_num', '?')}] {c.get('filename', 'Unknown')} (similarity: {c.get('score', 0):.3f})\n"

    if fmt == "md":
        content = f"# Query: {query}\n\n## Answer\n\n{answer}\n{sources_text}"
        return Response(
            content,
            mimetype="text/markdown",
            headers={"Content-Disposition": f"attachment;filename=doxium-answer-{task_id[:8]}.md"}
        )

    elif fmt == "txt":
        content = f"Query: {query}\n\nAnswer:\n{answer}\n{sources_text}"
        return Response(
            content,
            mimetype="text/plain",
            headers={"Content-Disposition": f"attachment;filename=doxium-answer-{task_id[:8]}.txt"}
        )

    elif fmt == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(f"Query: {query}", level=1)
        doc.add_heading("Answer", level=2)
        doc.add_paragraph(answer)
        doc.add_heading("Sources", level=2)
        for c in chunks:
            doc.add_paragraph(f"[{c.get('source_num', '?')}] {c.get('filename', 'Unknown')} (similarity: {c.get('score', 0):.3f})")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
                         download_name=f"doxium-answer-{task_id[:8]}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif fmt == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        buf = io.BytesIO()
        pdf_doc = SimpleDocTemplate(buf, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph(f"Query: {query}", styles["Title"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Answer", styles["Heading2"]))
        story.append(Paragraph(answer.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Sources", styles["Heading2"]))
        for c in chunks:
            story.append(Paragraph(f"[{c.get('source_num', '?')}] {c.get('filename', 'Unknown')} (similarity: {c.get('score', 0):.3f})", styles["BodyText"]))
        pdf_doc.build(story)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
                         download_name=f"doxium-answer-{task_id[:8]}.pdf",
                         mimetype="application/pdf")

    flash("Unsupported export format.", "danger")
    return redirect(url_for("main.dashboard"))


# -----------------------
# Legal Routes
# -----------------------
@main.route("/terms")
def terms():
    return render_template("terms.html")

@main.route("/privacy")
def privacy():
    return render_template("privacy.html")

@main.route("/refund")
def refund():
    return render_template("refund.html")


# -----------------------
# Admin Routes
# -----------------------
@main.route("/admin")
@login_required
def admin():
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.dashboard"))
    users = User.query.filter_by(is_pilot=True).all()
    return render_template("admin_pilots.html", users=users)


@main.route("/reset_index/<int:user_id>", methods=["POST"])
@csrf.exempt
@login_required
def reset_index(user_id):
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.admin"))
    try:
        embedder = EmbeddingGenerator()
        faiss_index = FaissIndex(dim=embedder.get_dimension(), user_id=user_id)
        faiss_index.rebuild_index_from_chunks()
        flash(f"Index reset for user {user_id}.", "success")
    except Exception as e:
        flash(f"Error resetting index: {str(e)}", "danger")
    return redirect(url_for("main.admin"))


@main.route("/toggle-pilot/<int:user_id>", methods=["POST"])
@csrf.exempt
@login_required
def toggle_pilot(user_id):
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.admin"))
    user = User.query.get_or_404(user_id)
    user.is_pilot = not user.is_pilot
    db.session.commit()
    flash(f"Pilot status for {user.username} updated.", "success")
    return redirect(url_for("main.admin"))


@main.route("/user-stats/<int:user_id>")
@login_required
def user_stats(user_id):
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.admin"))
    user = User.query.get_or_404(user_id)
    file_count = File.query.filter_by(user_id=user.id).count()
    has_space, available_mb, used_mb = check_storage_space(user.id)
    return render_template("admin_user_stats.html", user=user, file_count=file_count, used_mb=used_mb, available_mb=available_mb)


@main.route("/impersonate/<int:user_id>", methods=["POST"])
@csrf.exempt
@login_required
def impersonate(user_id):
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.admin"))
    user = User.query.get_or_404(user_id)
    from flask import session
    session["admin_id"] = current_user.id
    login_user(user)
    flash(f"You are now impersonating {user.username}.", "info")
    return redirect(url_for("main.dashboard"))


@main.route("/return-to-admin")
@login_required
def return_to_admin():
    from flask import session
    admin_id = session.get("admin_id")
    if admin_id:
        admin_user = User.query.get(admin_id)
        if admin_user and admin_user.is_admin:
            login_user(admin_user)
            session.pop("admin_id", None)
            flash("Returned to admin account.", "info")
            return redirect(url_for("main.admin"))
    flash("Unable to return to admin.", "danger")
    return redirect(url_for("main.dashboard"))


@main.route("/delete-user/<int:user_id>", methods=["POST"])
@csrf.exempt
@login_required
def delete_user(user_id):
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.admin"))
    user = User.query.get_or_404(user_id)
    for file in user.files:
        for chunk in file.chunks:
            db.session.delete(chunk)
        if os.path.exists(file.path):
            os.remove(file.path)
        db.session.delete(file)

    import shutil
    user_folder = os.path.join(current_app.config["UPLOAD_FOLDER"], str(user_id))
    if os.path.exists(user_folder):
        shutil.rmtree(user_folder)

    from .faiss_index import FAISS_BASE_PATH
    faiss_dir = os.path.join(FAISS_BASE_PATH, f"{user_id}.faiss")
    if os.path.exists(faiss_dir):
        shutil.rmtree(faiss_dir)

    db.session.delete(user)
    db.session.commit()
    flash("User and their data deleted.", "danger")
    return redirect(url_for("main.admin"))


@main.route("/system-stats")
@login_required
def system_stats():
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for("main.admin"))
    user_count = User.query.count()
    pilot_count = User.query.filter_by(is_pilot=True).count()
    file_count = File.query.count()
    chunk_count = Chunk.query.count()
    total_used = sum(f.size for f in File.query.all())
    total_used_mb = total_used / (1024 * 1024)
    return render_template("admin_stats.html", user_count=user_count, pilot_count=pilot_count,
                           file_count=file_count, chunk_count=chunk_count, total_used_mb=total_used_mb)


# -----------------------
# Error Handlers
# -----------------------
@main.errorhandler(404)
def page_not_found(e):
    return render_template("404.html"), 404


@main.errorhandler(500)
def internal_server_error(e):
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify(success=False, error="Internal server error"), 500
    return render_template("500.html"), 500
